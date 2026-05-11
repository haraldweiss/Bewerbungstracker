# SPDX-License-Identifier: AGPL-3.0-or-later
# © 2026 Harald Weiss
"""Export-Service für Cover Letters.

Wandelt das HTML-Content (mit data-confidence Attributen) in saubere PDF/DOCX
Dateien um. Confidence-Attribute werden vor dem Export entfernt — die sind nur
fürs Frontend-Coloring relevant, nicht für den finalen Bewerbungstext.

Dependencies werden lazy importiert, damit das Modul auch ohne reportlab/docx
ladbar ist (für Tests die Mocks nutzen).
"""
from __future__ import annotations
import io
import re
import logging
from datetime import datetime
from html import unescape
from typing import Optional

logger = logging.getLogger(__name__)

# Style-Konstanten (konsistent für PDF + DOCX)
FONT_NAME = 'Helvetica'  # PDF
DOCX_FONT = 'Calibri'    # DOCX
BODY_SIZE_PT = 11
LINE_SPACING = 1.15
PAGE_MARGIN_INCH = 1.0


def _strip_confidence_attributes(html: str) -> str:
    """Entfernt data-confidence und data-source Attribute aus HTML.

    Diese sind nur für Frontend-Coloring. Im finalen Bewerbungs-Dokument
    haben sie nichts zu suchen.
    """
    cleaned = re.sub(r'\s*data-confidence="[^"]*"', '', html)
    cleaned = re.sub(r'\s*data-source="[^"]*"', '', cleaned)
    return cleaned


def _extract_paragraphs(html: str) -> list[str]:
    """Holt nur den Text aus <p>...</p> Blöcken — Markup wird gestrippt."""
    paragraphs = re.findall(r'<p[^>]*>(.*?)</p>', html, re.DOTALL)
    out = []
    for raw in paragraphs:
        # Inline-HTML stripen (z.B. <br>, <strong>)
        text = re.sub(r'<[^>]+>', '', raw)
        text = unescape(text).strip()
        if text:
            out.append(text)
    return out


class ExportService:
    """Rendert Cover-Letter-HTML zu PDF oder DOCX."""

    def to_pdf(self, html_content: str, applicant_name: str, company_name: str,
               job_title: str, applicant_address: Optional[str] = None) -> bytes:
        """Generiert PDF aus HTML-Content. Returns: PDF-Bytes."""
        # Lazy import — reportlab evtl. nicht installiert in CI/Tests
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib.enums import TA_LEFT, TA_JUSTIFY
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

        clean = _strip_confidence_attributes(html_content)
        paragraphs = _extract_paragraphs(clean)

        buf = io.BytesIO()
        doc = SimpleDocTemplate(
            buf, pagesize=A4,
            topMargin=PAGE_MARGIN_INCH*inch, bottomMargin=PAGE_MARGIN_INCH*inch,
            leftMargin=PAGE_MARGIN_INCH*inch, rightMargin=PAGE_MARGIN_INCH*inch,
            title=f"Anschreiben - {company_name}",
            author=applicant_name,
        )

        styles = getSampleStyleSheet()
        elements = []

        # Header: Absender
        name_style = ParagraphStyle('Name', parent=styles['Normal'],
            fontName=FONT_NAME, fontSize=12, spaceAfter=4, alignment=TA_LEFT)
        elements.append(Paragraph(f"<b>{applicant_name}</b>", name_style))
        if applicant_address:
            addr_style = ParagraphStyle('Addr', parent=styles['Normal'],
                fontName=FONT_NAME, fontSize=10, spaceAfter=12, alignment=TA_LEFT)
            for line in applicant_address.splitlines():
                if line.strip():
                    elements.append(Paragraph(line.strip(), addr_style))
        elements.append(Spacer(1, 0.2*inch))

        # Datum
        date_style = ParagraphStyle('Date', parent=styles['Normal'],
            fontName=FONT_NAME, fontSize=10, alignment=TA_LEFT, spaceAfter=18)
        elements.append(Paragraph(datetime.now().strftime('%d. %B %Y'), date_style))

        # Subject
        subj_style = ParagraphStyle('Subj', parent=styles['Normal'],
            fontName=FONT_NAME, fontSize=11, alignment=TA_LEFT, spaceAfter=18)
        elements.append(Paragraph(f"<b>Bewerbung als {job_title}</b>", subj_style))

        # Body
        body_style = ParagraphStyle('Body', parent=styles['Normal'],
            fontName=FONT_NAME, fontSize=BODY_SIZE_PT,
            spaceAfter=12, leading=14, alignment=TA_JUSTIFY)
        for para in paragraphs:
            elements.append(Paragraph(para, body_style))

        # Signatur-Platzhalter
        elements.append(Spacer(1, 0.3*inch))
        sig_style = ParagraphStyle('Sig', parent=styles['Normal'],
            fontName=FONT_NAME, fontSize=11, alignment=TA_LEFT)
        elements.append(Paragraph("Mit freundlichen Grüßen", sig_style))
        elements.append(Spacer(1, 0.4*inch))
        elements.append(Paragraph(applicant_name, sig_style))

        doc.build(elements)
        return buf.getvalue()

    def to_docx(self, html_content: str, applicant_name: str, company_name: str,
                job_title: str, applicant_address: Optional[str] = None) -> bytes:
        """Generiert DOCX aus HTML-Content. Returns: DOCX-Bytes (ZIP-Format)."""
        from docx import Document
        from docx.shared import Pt, Inches

        clean = _strip_confidence_attributes(html_content)
        paragraphs = _extract_paragraphs(clean)

        doc = Document()
        doc.core_properties.title = f"Anschreiben - {company_name}"
        doc.core_properties.author = applicant_name

        # Margins
        for section in doc.sections:
            section.top_margin = Inches(PAGE_MARGIN_INCH)
            section.bottom_margin = Inches(PAGE_MARGIN_INCH)
            section.left_margin = Inches(PAGE_MARGIN_INCH)
            section.right_margin = Inches(PAGE_MARGIN_INCH)

        # Default-Style auf Calibri 11pt setzen
        style = doc.styles['Normal']
        style.font.name = DOCX_FONT
        style.font.size = Pt(BODY_SIZE_PT)

        # Header: Absender
        header_p = doc.add_paragraph()
        run = header_p.add_run(applicant_name)
        run.bold = True
        run.font.size = Pt(12)

        if applicant_address:
            for line in applicant_address.splitlines():
                if line.strip():
                    p = doc.add_paragraph(line.strip())
                    p.runs[0].font.size = Pt(10)

        doc.add_paragraph()  # spacer

        # Datum
        date_p = doc.add_paragraph(datetime.now().strftime('%d. %B %Y'))
        date_p.runs[0].font.size = Pt(10)
        doc.add_paragraph()

        # Subject
        subj_p = doc.add_paragraph()
        subj_run = subj_p.add_run(f"Bewerbung als {job_title}")
        subj_run.bold = True
        doc.add_paragraph()

        # Body
        for para in paragraphs:
            p = doc.add_paragraph(para)
            p.paragraph_format.line_spacing = LINE_SPACING
            p.paragraph_format.space_after = Pt(8)

        # Signatur
        doc.add_paragraph()
        doc.add_paragraph("Mit freundlichen Grüßen")
        doc.add_paragraph()
        doc.add_paragraph(applicant_name)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()
